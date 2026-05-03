"""
services/resume_service.py – Resume text extraction from PDF/DOCX.
Per project rules: extract text, store in DB, discard file. Never persist to disk.
"""
import io
import logging
import re
from typing import Optional, Dict

logger = logging.getLogger(__name__)


async def extract_resume_text(file_bytes: bytes, filename: str) -> Optional[dict]:
    """
    Extract plain text and links from PDF or DOCX bytes.
    Returns: {"text": str, "links": {"linkedin": str, "github": str, "portfolio": str, "others": list[str]}}
    """
    fname = (filename or "").lower()
    text = None
    try:
        if fname.endswith(".pdf"):
            text = _extract_pdf(file_bytes)
        elif fname.endswith(".docx"):
            text = _extract_docx(file_bytes)
        elif fname.endswith(".doc"):
            text = _extract_doc_fallback(file_bytes)
        elif fname.endswith(".txt"):
            text = file_bytes.decode("utf-8", errors="ignore")
        
        # Try MarkItDown if installed for better markdown-style extraction
        try:
            from markitdown import MarkItDown
            mid = MarkItDown()
            # MarkItDown usually takes a file path, we can use a temporary buffer
            import tempfile
            import os
            with tempfile.NamedTemporaryFile(suffix=os.path.splitext(fname)[1], delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            
            mid_result = mid.convert(tmp_path)
            if mid_result and mid_result.text_content:
                text = mid_result.text_content
            
            os.unlink(tmp_path)
        except (ImportError, Exception) as e:
            logger.debug(f"MarkItDown extraction skipped or failed: {e}")

        if not text:
            return None

        links = _extract_links(text)
        return {
            "text": text,
            "links": links
        }
    except Exception as e:
        logger.error(f"Resume extraction failed for {filename}: {e}", exc_info=True)
        return None


def _extract_links(text: str) -> Dict:
    """Extract LinkedIn, GitHub, and other links from text."""
    links = {
        "linkedin": None,
        "github": None,
        "portfolio": None,
        "others": []
    }
    
    # Generic URL regex
    url_pattern = r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[/\w\.-]*'
    all_urls = re.findall(url_pattern, text)
    
    for url in all_urls:
        url_lower = url.lower()
        if "linkedin.com/in/" in url_lower:
            links["linkedin"] = url
        elif "github.com/" in url_lower:
            links["github"] = url
        elif any(domain in url_lower for domain in ["behance.net", "dribbble.com", "portfolio", "personal-site"]):
            links["portfolio"] = url
        else:
            if url not in links["others"]:
                links["others"].append(url)
                
    return links


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
    except ImportError:
        logger.warning("pdfplumber not installed — trying PyPDF2 fallback")
        return _extract_pdf_pypdf(file_bytes)
    except Exception as e:
        logger.warning(f"PDF extraction with pdfplumber failed: {e}")
        return _extract_pdf_pypdf(file_bytes)


def _extract_pdf_pypdf(file_bytes: bytes) -> str:
    """Fallback PDF extraction using PyPDF2."""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = [p.extract_text() or "" for p in reader.pages]
        return "\n\n".join(p for p in pages if p.strip())
    except ImportError:
        raise RuntimeError("No PDF extraction library available. Install pdfplumber.")


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table text
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def _extract_doc_fallback(file_bytes: bytes) -> str:
    """Best-effort .doc extraction — strip binary, keep printable text."""
    try:
        text = file_bytes.decode("latin-1", errors="ignore")
        # Remove non-printable chars except newlines/tabs
        import re
        clean = re.sub(r"[^\x20-\x7e\n\t]", " ", text)
        # Collapse whitespace runs
        clean = re.sub(r"  +", " ", clean)
        return clean.strip()
    except Exception:
        return ""


def validate_resume_file(filename: str, size_bytes: int) -> Optional[str]:
    """
    Returns an error message if the file is invalid, else None.
    """
    MAX_SIZE = 5 * 1024 * 1024  # 5 MB
    ALLOWED_EXT = {".pdf", ".docx", ".doc", ".txt"}

    if size_bytes > MAX_SIZE:
        return "File too large. Maximum size is 5 MB."

    import os
    ext = os.path.splitext(filename.lower())[1]
    if ext not in ALLOWED_EXT:
        return f"Invalid file type '{ext}'. Please upload a PDF, DOCX, or TXT file."

    return None
